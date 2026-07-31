from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = "/Users/kaviya/products/plumtale/MoonBell_Narrow_Paid_Beta_Plan.docx"
BLUE = RGBColor(46,116,181); DARK = RGBColor(31,77,120); INK = RGBColor(30,30,30); MUTED = RGBColor(90,90,90)

doc = Document()
sec = doc.sections[0]
for attr in ("top_margin","bottom_margin","left_margin","right_margin"):
    setattr(sec, attr, Inches(1))
sec.header_distance = Inches(0.492); sec.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal._element.rPr.rFonts.set(qn("w:ascii"),"Calibri"); normal._element.rPr.rFonts.set(qn("w:hAnsi"),"Calibri")
normal.font.size = Pt(11); normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.25
for name,size,color,before,after in [("Heading 1",16,BLUE,18,10),("Heading 2",13,BLUE,14,7),("Heading 3",12,DARK,10,5)]:
    st=doc.styles[name]; st.font.name="Calibri"; st._element.rPr.rFonts.set(qn("w:ascii"),"Calibri"); st._element.rPr.rFonts.set(qn("w:hAnsi"),"Calibri")
    st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=color
    st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn("w:shd"))
    if shd is None: shd=OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"),fill)

def set_cell(cell,text,bold=False,size=8.7):
    cell.text=""; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(2); p.paragraph_format.line_spacing=1.05
    r=p.add_run(str(text)); r.font.name="Calibri"; r._element.rPr.rFonts.set(qn("w:ascii"),"Calibri"); r._element.rPr.rFonts.set(qn("w:hAnsi"),"Calibri")
    r.font.size=Pt(size); r.bold=bold; r.font.color.rgb=DARK if bold else INK
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tcPr=cell._tc.get_or_add_tcPr(); tcMar=OxmlElement("w:tcMar")
    for m,v in [("top",80),("start",120),("bottom",80),("end",120)]:
        n=OxmlElement("w:"+m); n.set(qn("w:w"),str(v)); n.set(qn("w:type"),"dxa"); tcMar.append(n)
    tcPr.append(tcMar)

def table(headers, rows, widths=None, size=8.7):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.autofit=False
    for i,h in enumerate(headers):
        set_cell(t.rows[0].cells[i],h,True,size); shade(t.rows[0].cells[i],"E8EEF5")
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,v in enumerate(row):
            set_cell(cells[i],v,False,size)
            if ri%2: shade(cells[i],"FAFBFC")
    tblPr=t._tbl.tblPr; w=tblPr.find(qn("w:tblW"))
    if w is None: w=OxmlElement("w:tblW"); tblPr.append(w)
    w.set(qn("w:w"),"9360"); w.set(qn("w:type"),"dxa")
    ind=OxmlElement("w:tblInd"); ind.set(qn("w:w"),"120"); ind.set(qn("w:type"),"dxa"); tblPr.append(ind)
    grid=t._tbl.tblGrid
    for c in list(grid): grid.remove(c)
    dx=[round(x/6.5*9360) for x in widths] if widths else [round(9360/len(headers))]*len(headers)
    for x in dx:
        c=OxmlElement("w:gridCol"); c.set(qn("w:w"),str(x)); grid.append(c)
    for row in t.rows:
        for i,c in enumerate(row.cells):
            tcPr=c._tc.get_or_add_tcPr(); tcW=OxmlElement("w:tcW"); tcW.set(qn("w:w"),str(dx[i])); tcW.set(qn("w:type"),"dxa"); tcPr.append(tcW)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return t

def para(text="", bold=False, italic=False, color=INK, size=11, align=None):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.25
    if align is not None: p.alignment=align
    r=p.add_run(text); r.font.name="Calibri"; r._element.rPr.rFonts.set(qn("w:ascii"),"Calibri"); r._element.rPr.rFonts.set(qn("w:hAnsi"),"Calibri")
    r.font.size=Pt(size); r.bold=bold; r.italic=italic; r.font.color.rgb=color
    return p

def bullet(text):
    p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.left_indent=Inches(.375); p.paragraph_format.first_line_indent=Inches(-.188); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.25; p.add_run(text)

def num(text):
    p=doc.add_paragraph(style="List Number"); p.paragraph_format.left_indent=Inches(.375); p.paragraph_format.first_line_indent=Inches(-.188); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.25; p.add_run(text)

# furniture
h=sec.header.paragraphs[0]; h.text="MOONBELL  /  BETA OPERATING PLAN"; h.alignment=WD_ALIGN_PARAGRAPH.RIGHT
for r in h.runs: r.font.size=Pt(8.5); r.font.color.rgb=MUTED
f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER; rr=f.add_run("Confidential working document  •  31 July 2026"); rr.font.size=Pt(8); rr.font.color.rgb=MUTED

para("MOONBELL",True,False,BLUE,12,WD_ALIGN_PARAGRAPH.CENTER)
para("Narrow Paid Beta Execution Plan",True,False,INK,28,WD_ALIGN_PARAGRAPH.CENTER)
para("A deliberately capped operating design to validate demand, delivery quality, and unit economics",False,False,MUTED,14,WD_ALIGN_PARAGRAPH.CENTER)
para("Decision: do not accept paid orders until the release gates are complete.",True,False,RGBColor(155,28,28),11,WD_ALIGN_PARAGRAPH.CENTER)
table(["Scope","Default product","Beta cap","Price test"],[["India-first physical beta","English birthday hardcover","75 stranger-paid orders","₹1,299 vs ₹1,499"]],[1.7,1.8,1.5,1.5],9)
para("This is a validation plan, not a product roadmap. Every unproven feature is intentionally excluded.",True,False,MUTED,10.5,WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

doc.add_heading("Executive recommendation",1)
para("Do not launch the current broad MoonBell offer. Launch one invite-only birthday hardcover to a maximum of 75 stranger customers, only after provider, legal, print, payment, and fulfilment gates are closed.")
para("Potentially fatal gate: Google Cloud’s current Generative AI terms prohibit use in an online service directed toward or likely to be accessed by under-18s. Obtain authoritative written clearance for the parent-operated, child-output workflow or use a contractually suitable provider before even generating external previews.")
for x in ["[MUST] before accepting money","[MANUAL] founder-run during beta","[OFF] explicitly disabled","[LATER] after beta success"]: bullet(x)

doc.add_heading("1. The narrowed beta",1)
table(["Decision","Exact beta"],[
("Buyer","English-speaking parent/legal guardian of a child aged 3–7; own child"),
("Occasion","Birthday at least 21 days away"),
("Product","One English Birthday Adventure hardcover; 8×8 inches; casebound; matte laminated"),
("Pages","Fixed 20 interior pages: 12 illustrated story pages plus title, dedication, copyright, fixed end matter"),
("Inputs","Nickname, age band, optional stated pronouns, skin tone, hair length/texture/style, glasses, up to three interests, one short dedication"),
("Forbidden","Legal name, exact birthdate, school, location, health data, free-form prompt, photo"),
("Preview","Cover plus first three illustrated pages; nine pages locked; no full text or PDF"),
("Correction","One founder-reviewed objective correction; no open-ended regeneration"),
("Price","₹1,299 versus ₹1,499, including eligible shipping and applicable tax; one price per visitor"),
("Geography","Allowlisted urban pincodes in Bengaluru, Mumbai, and Delhi NCR"),
("Capacity","75 total; 25 per metro; 3/day and 15/week; pause above five books awaiting QC"),
("Promise","Dispatch within 7 business days; delivery within 14 business days"),
("Refund/replacement","Full refund before print release; one free reprint/reship for MoonBell error, material defect, damage, missing pages, wrong personalization, or carrier loss"),
], [1.55,4.95],8.6)
para("Do not test ₹999 unless verified quotes radically reduce repeatable cost. At current assumptions it is close to zero contribution before CAC.",True,False,MUTED,10.5)

doc.add_heading("2. Explicit beta scope",1)
doc.add_heading("Included — [MUST]",2)
for x in ["Invite-gated access; parent/guardian-only ordering; pincode validation.","Three-step intake; transactional email; versioned consent; cover-plus-three-page preview.","One server-controlled price; prepaid Razorpay; webhook-confirmed payment; one hardcover.","Human story/illustration/print QC; physical inspection; tracking; support; replacement/refund SOP.","Per-order ledger and PII-free funnel analytics."]: bullet(x)
doc.add_heading("Excluded — [OFF]",2)
for x in ["Standalone PDF, audio, photo upload, non-birthday themes, custom prompts, other languages, paperback, format/page-count choices.","COD, delivery outside allowlisted pincodes, non-guardian gifting, multiple books, public sharing, automated whole-story regeneration.","Subscriptions, referrals, coupons, bundles, schools, fairs, marketplaces, newsletters, birthday reminders, coming-soon products."]: bullet(x)
doc.add_heading("Manual operations — [MANUAL]",2)
for x in ["Invites, pincode checks, capacity, corrections, editorial/illustration/print QC, printer upload, invoices, packing, courier, tracking, NDR/RTO, claims, refunds, reprints, support, interviews, ledger."]: bullet(x)
doc.add_heading("Automated requirements — [MUST]",2)
for x in ["Price/cap and one payable order per book; signature/amount/currency/event verification; duplicate/out-of-order webhook handling.","Captured payment as fulfilment trigger; refund/cancellation stopping print release; immutable release timestamp and file hash; deletion guard.","Photo rejection; fail-closed moderation; parent-scoped access; signed links; server-side locked pages; PII-free analytics and alerts."]: bullet(x)

doc.add_heading("3. Pre-beta blockers",1)
blockers=[
("Provider clearance","Google/Vertex age restriction is potentially fatal. Obtain written authorization or change provider.","Founder + technology counsel","Signed exception/authoritative interpretation","PV/PAY/PR"),
("Seller identity","MoonBell must have one accountable legal seller across site, Razorpay, contracts and invoices.","Founder + CA/CS/lawyer","Constitution/PAN/bank/address/contact aligned","PV/PAY"),
("GST/invoice","Printed books may be nil-rated, but the personalized physical supply and included elements need classification.","E-commerce CA","Signed HSN/SAC, registration and invoice/bill memo","PAY/SHIP"),
("Legal pages","Current pages still describe digital-only launch and contain TODO identity fields.","Counsel + founder","No TODOs; counsel sign-off; copy matches checkout/SOP","PV/PAY/SHIP"),
("Consent/privacy","Map local storage, AI, email, printer, courier, Razorpay, retention, deletion exceptions.","Privacy counsel + engineer","Signed map; versioned consent; deletion tests","PV/PAY"),
("Photo off","A hidden UI is not enough if direct upload/egress remains possible.","Engineer","UI absent; API rejects; no bucket writes/egress","PV"),
("AI/IP disclosure","Do not promise exact likeness, exclusive copyright, or hand illustration.","IP counsel","Rights matrix and pre-purchase disclosure","PAY/PR"),
("Printer/RFQ/contract","Three comparable quotes plus confidentiality, deletion, subcontractor, breach and reprint controls.","Founder + printer + counsel","RFQs; signed SOW/NDA/data terms","PR"),
("Interior PDF","Current assembler is trim-only; prove bleed, safe area, gutter, fonts, color, page order and PDF standard.","Prepress + engineer + printer","Passed preflight report","PR"),
("Casewrap/spine","Hardcover requires front/back/spine/hinges/turn-ins and a page-count-specific template.","Prepress + printer","Printer template and approved export","PR"),
("Resolution","Low-resolution AI art can fail at physical size.","Engineer + prepress","300 PPI target; 250 PPI minimum at placed size","PR"),
("Physical proof","Screen review cannot prove binding, colour, crop, durability, or packaging.","Founder/QC + printer","12 proofs; signed acceptance sheet","PR/SHIP"),
("Razorpay live","Test mode does not prove capture, refund, settlement, descriptor, or live webhook.","Founder/finance + engineer","Low-value live capture, refund, settlement","PAY"),
("Payment idempotency","Current order endpoint can create new orders on retries/double-clicks.","Engineer","Concurrency/retry tests prove one payable order","PAY/PR"),
("Webhook/reconciliation","Duplicate/out-of-order events and forged callbacks must not trigger fulfilment.","Engineer + finance","Replay suite; unmatched-payment alert; daily reconciliation","PAY/PR"),
("Refund/print guard","Refund before print release must hold/cancel fulfilment atomically.","Engineer + ops","Concurrent refund/release tests","PAY/PR"),
("Deletion guard","Paid/printed/shipped orders cannot be deleted mid-fulfilment.","Engineer + privacy counsel + CA","Deferred-erasure tests and retention matrix","PAY/PR/SHIP"),
("QC/reprint","Every paid book needs a signed human checklist and one linked replacement case.","Founder/QC","Checklist, hash, reviewer, timestamp, dry-run reprint","PR/SHIP"),
("Courier/support","Allowlist, rate card, NDR/RTO/loss/damage claims, support and grievance process must exist.","Founder/ops + courier","Test shipments; complaint register; visible contacts","SHIP/PAY"),
("Incident/logs","CERT-In readiness requires a point of contact, 6-hour decision path and 180-day logs.","Security lead + cyber counsel","PoC, logs, alerts, timed tabletop","PV/PAY"),
]
table(["Gate","Why","Owner","Evidence","Blocks"],blockers,[1.2,2.25,1.1,1.55,.7],7.5)

doc.add_heading("4. Printing validation plan",1)
for x in ["Approach exactly three printers; prefer one close enough for founder inspection/collection.","Send one RFQ covering 8×8 trim, fixed 20 pages, casebound construction, paper/board targets, PDF/X, ICC profile, bleed, safe area, gutter, spine formula, variable files, MOQ, proofs, SLA, packaging, tax, reprints, and subcontractors.","Prepare the test files yourself: full interior and casewrap; full bleed/non-bleed; dark scenes; Indian skin tones; fine detail; small type; cross-spreads; edge-adjacent elements.","First proof round: six copies, two per printer. One per printer to the founder; one to a test address in one of the three beta metros.","Select the printer, sign the SOW, then print six final proofs: two distinct addresses in each beta metro.","Inspect trim, bleed, margins, gutter, text, pixelation, colour, binding, hinge, lamination, spine alignment, page order, opening, glue, board, and package protection.","Run one-metre face/edge/corner drops, 10 kg compression, shake, light splash, rub/scuff, and open/reseal tests on founder copies."]: num(x)
para("Acceptance: zero preflight errors; 300 PPI target and 250 PPI minimum; zero major final-proof defects; 6/6 correct personalization and page order; 6/6 packages intact; no undocumented printer changes.",True,False,DARK,10.5)
para("Reject wrong names, missing/misordered pages, cropped text, white bleed edges, obvious pixelation, severe colour cast, binding failure, warped boards, damaged packages, or unsafe content.",False,False,RGBColor(155,28,28),10.5)

doc.add_heading("5. Beta order lifecycle",1)
ops=[
("Preview","PREVIEW_GENERATING → PREVIEW_READY","System; founder exceptions","Intake, consent, model, cost; provider/moderation/email failures","Ready link or no-charge failure","Hybrid"),
("Payment","CHECKOUT_STARTED → ORDER_CREATED","Customer/system","Frozen price, pincode, address, cap; duplicate/invalid failure","Price, delivery, cancellation cutoff before Razorpay","Automated"),
("Verification","PAYMENT_CAPTURED / REVIEW","Webhook + finance","IDs, amount, currency; mismatch/replay/late auth","Confirm only after capture","Hybrid"),
("Generation","PAID_GENERATING → PRINT_MASTER_READY","Pipeline","Remaining pages, cost, moderation","Progress only if SLA risk","Automated"),
("QC","QC_PENDING → PASS / HOLD","Founder/QC","Story, name, age, appearance, pages, images, print checklist","No message unless delayed","Manual"),
("Correction","CUSTOMER_ACTION_REQUIRED","Founder","Issue, correction, approval, cost","One specific question and new ETA","Manual"),
("Print release","RELEASED_TO_PRINT","Founder + system guard","Paid/not-refunded, PDF/casewrap hash, reviewer, timestamp","Changes close after release","Hybrid"),
("Production","PRINTING → PRINTED_QC","Printer + founder","Job ID, preflight, print date, physical pass","Notify only if estimate changes","Manual"),
("Shipment","SHIPMENT_CREATED → SHIPPED","Founder","AWB, courier, address, weight/value","Tracking and window","Manual"),
("Delivery","IN_TRANSIT / NDR / DELIVERED","Founder","Scans, NDR reason, proof of delivery","Same-day NDR contact","Manual"),
("Exception","REPRINT / RESHIP / RTO / LOST / REFUND","Founder","Ticket, evidence, fault, cost, claim, remedy","Exact remedy and new date","Manual + guard"),
("Feedback","FEEDBACK_PENDING → CLOSED","Founder","Story/physical score, interview, referral","One request 24–72h after delivery","Manual"),
]
table(["Step","Status","Owner","Data/failure","Customer message","Mode"],ops,[.75,1.05,.85,2.05,1.15,.65],7.4)

doc.add_heading("6. Unit economics",1)
table(["Ledger group","Record"],[
("Cohort","Order/book/preview IDs; stranger; city; source; invite; price arm"),
("Funnel","Preview, payment, release, dispatch, delivery timestamps; outcome"),
("Revenue/tax","Price, discount, captured cash, refund, output tax, net revenue"),
("AI","Preview text/images/moderation/failed attempts/correction; remaining-book AI"),
("Physical","Print, setup/preflight, packaging, forward/reverse shipping"),
("Payment","Razorpay fee and GST on fee"),
("Exceptions","Reprint, repack, reship, RTO, loss, damage, refund"),
("Labour","QC, fulfilment, support, grievance, acquisition minutes at ≥₹300/hour"),
("Marketing","Placement fees, samples, commissions, attributable spend"),
("Quality","Defect, contacts, story score, physical score, satisfaction, referral"),
],[1.35,5.15],8.5)
for x in ["Net revenue = captured cash − refunds − output tax.","Preview burden/order = all preview, failed-generation and pre-payment correction cost ÷ webhook-confirmed stranger orders.","Fully loaded CAC = cash acquisition spend + samples + acquisition labour ÷ webhook-confirmed stranger orders.","Contribution before CAC = net revenue − preview burden − paid-book AI − print − packaging − shipping/RTO/reprint reserve − Razorpay fee − fulfilment/QC/support labour − repeatable infrastructure.","Contribution after CAC = contribution before CAC − fully loaded CAC."]: bullet(x)
table(["Price","Before CAC, nil-rated","After ₹150 CAC","Before CAC with 18% tax","After CAC with 18% tax"],[["₹1,299","~₹331","~₹181","~₹133","~−₹17"],["₹1,499","~₹527","~₹377","~₹298","~₹148"]],[1.0,1.35,1.2,1.35,1.25],8.2)
para("At 25% preview conversion, preview burden falls from roughly ₹280 to ₹168/order. Do not test ₹999 unless verified quotes radically reduce repeatable cost.",True,False,MUTED,10.5)

doc.add_heading("7. Paid beta experiment",1)
for x in ["Stage 1: up to 60 stranger-paid orders, 30 at each price, ten per price per metro, alternating matched blocks, maximum 200 valid viewed previews per arm.","Stage 2: conditional final 15 orders, five per metro at the winning price. If neither price qualifies, stop at 60 or fewer.","Allowed recruitment: existing waitlist strangers, admin-approved parent communities, and up to three paid local parent-community/micro-creator placements.","No free books, cashback, hidden coupons, unrecorded shipping subsidy, free samples omitted from CAC, or zero-valued founder labour.","Ask after delivery: what made you buy, what nearly stopped you, whether preview was sufficient, strongest/weakest page, consistency, price willingness, and birthday timing.","Prioritize behaviour: captured payment, contribution per preview, high-price conversion, abandonment, corrections, refunds/reprints, delivery, sharing, paid referrals, then survey scores."]: bullet(x)

doc.add_heading("8. Success / change / stop thresholds",1)
table(["Metric","Continue","Change","Stop"],[
("Preview→paid ₹1,299","≥20%","16–19.9%","<16%"),
("Preview→paid ₹1,499","≥15%","10–14.9%","<10%"),
("Contribution before CAC","≥₹350","₹150–349","<₹150"),
("Contribution after CAC","≥₹150 and ≥12%","₹0–149","Negative"),
("CAC","≤₹150","₹151–250",">₹250"),
("On-time delivery","≥95%","90–94.9%","<90%"),
("Defect/damage/reprint/refund","≤5%","5.1–10%",">10%"),
("Support time","Median ≤15m; P90 ≤30m","Median 16–30m","Median >30m"),
("Satisfaction","≥85% rate 8–10","70–84%","<70%"),
("Story quality","≥4.3/5","3.8–4.29","<3.8"),
("Physical quality","≥4.5/5","4.0–4.49","<4.0"),
("Paid referral within 30d","≥8%","3–7.9%","<3%"),
],[2.0,1.5,1.5,1.5],8.1)
para("Hard stops override averages: duplicate charge/print, cross-customer exposure, unsafe content, unresolved provider/privacy breach, repeated serious printer defects, QC backlog above five, or delivery failure cluster.",True,False,RGBColor(155,28,28),10.5)

doc.add_heading("9. Do not build before beta success",1)
for x in ["More themes/occasions; languages; subscriptions; educational products; COD; mobile app; printer/courier integrations; large influencer campaigns; broad ads; schools/fairs/marketplaces; premium packaging; new formats; photo upload; complex dashboards; self-service editing/refunds; referral engineering; loyalty/bundles; audio/digital tiers; public sharing."]: bullet(x)

doc.add_heading("10. Dependency-ordered execution",1)
table(["Phase","Deliverables","Exit criteria","Do not begin"],[
("1 Legal/provider","Provider clearance, seller, CA tax, Legal Metrology, privacy/consent, policies, AI/IP, photo-off","Written clearance; counsel/CA sign-off; no TODOs; negative photo test","External previews, public acquisition, provider processing"),
("2 Printer/proof","Three RFQs, contract, preflight, casewrap, 12 proofs, packaging/delivery tests","Winning printer; zero major final-proof defects; fixed SLA","Paid orders or delivery claims"),
("3 Payment/fulfilment","Live Razorpay, idempotency, webhook, refund/deletion guard, invoice, courier, support, logs, analytics","Live capture/refund/settlement and replay/concurrency tests","Live buy button"),
("4 Dry runs","Six scenarios: three deliveries, duplicate/replay, refund-before-print, damaged/reprint/NDR","Zero duplicate charge/print; reconciled payments; lanes tested","Stranger recruitment"),
("5 Paid beta","60 price-test orders, conditional 15 winner orders, weekly reviews","Terminal outcomes; no hard stop; thresholds calculated","New product/geography/theme/channel"),
("6 Decision","Cohort P&L, price/quality/delivery analysis, interviews, 30-day referral","Written go/change/stop decision","Scale spend or roadmap expansion"),
],[1.0,2.3,1.8,1.4],7.8)

doc.add_heading("11. Landing page conversion redesign",1)
para("The local mobile render is approximately 15,262 pixels tall, with the main CTA below the first 390×844 viewport. It has roughly 15 /create links, eight occasions, future products, newsletters, and an invite wall. It proves illustration quality but not physical-product quality.")
table(["Element","Exact direction"],[
("Header","Logo; Sample; FAQ; one CTA. Remove account/books links for cold visitors."),
("Hero","Birthday-specific headline, exact assigned price/spec/timeline, CTA above fold, real hardcover photo."),
("Hero copy","A hardcover birthday storybook starring your child. Tell us their nickname, appearance and favourite things. See the cover and first 3 illustrated pages free."),
("CTA","Create my free preview. Secondary: See the sample book. Microcopy: About 3 minutes · no child photo · no payment until preview."),
("Trust row","Preview before payment; no child photo collected; human-reviewed before print; print damage replaced."),
("Product block","Real closed/open/casewrap/paper/packaging photos; exact size, pages, binding, price, shipping, timeline, replacement."),
("Story proof","Three sequential sample pages with readable prose, age-matching, story arc, interest integration, and consistency."),
("Personalization","One before/after showing nickname, appearance, and interests changing story/art; clarify original illustration, not exact portrait."),
("Process","Four steps: describe, preview, approve/pay, review/print/deliver."),
("Trust","Seller, privacy, AI role, Razorpay, human QC, serviceability, replacement, support. No placeholders or fake testimonials."),
("Pricing","One assigned price, shipping included, no coupon, no strike-through anchor."),
("FAQ","Preview, likeness, AI/data, arrival, correction, defects, cancellation, ordering eligibility."),
("Mobile","CTA in first viewport; sticky CTA after hero; single-column form; lazy-load below fold; no popup/video/chat."),
],[1.35,5.15],8.5)
para("Show cover plus three pages only; lock nine; do not expose full text through the API. Collect transactional email before generation without forcing a confirmation detour. Validate pincode before Razorpay. Keep checkout guest-first and minimal.",True,False,DARK,10.5)
para("Minimum analytics: landing_view, price_arm_assigned, CTA clicks, scroll depth, engaged time, sample views, preview start/completion/failure, intake steps/errors, email, consent, correction, checkout, pincode, payment initiated/failed/dismissed, webhook-captured purchase, print/shipping/delivery/refund/reprint, feedback, and paid referral. Never send child data to analytics.",False,False,INK,10.5)
para("First post-beta tests: physical hero shot; birthday headline; full product line; CTA copy; sample-versus-product order; trust placement; sticky CTA timing; email timing; two versus three preview pages; verified testimonial placement.",False,False,INK,10.5)

doc.add_heading("Final decision pack",1)
doc.add_heading("Exact beta",2)
para("One English 8×8-inch casebound Birthday Adventure hardcover for children aged 3–7; fixed 20-page interior with 12 illustrated story pages; no photo or digital product; cover plus three-page preview; ₹1,299 versus ₹1,499; three-metro allowlist; 75 stranger-paid orders maximum; founder QC and fulfilment; delivery within 14 business days.")
doc.add_heading("Top five blockers",2)
for x in ["Google/Vertex contractual ambiguity.","Unproven print-ready interior, casewrap, and physical quality.","Unfinished seller, tax, invoice, and legal-page position.","Payment idempotency plus refund/deletion/print-release safety.","Unproven courier exceptions, founder fulfilment, and actual unit cost."]: num(x)
doc.add_heading("Go / change / stop",2)
for x in ["GO only when provider/legal gates are clear, one price clears conversion and margin, contribution after CAC is ≥₹150/order, on-time delivery ≥95%, serious failure ≤5%, and story/physical quality pass.","CHANGE when there is no safety/legal hard stop and contribution is positive but below target; change one bounded variable and run a separately labelled cohort.","STOP/PIVOT when provider authorization is unresolved, both prices lose after CAC, quality/delivery remains unacceptable, or any recurring child-safety, cross-customer, or payment-integrity failure occurs."]: bullet(x)
doc.add_heading("First ten actions",2)
for x in ["Freeze the beta specification and remove ₹999 from the planned test.","Send the detailed workflow request to Google Cloud.","Archive/review OpenAI agreement, DPA, model status, and consent.","Name the legal seller and engage counsel and CA.","Obtain GST, invoice, and Legal Metrology treatment.","Send the same RFQ to exactly three printers.","Obtain printer templates and sign confidentiality/SOW terms.","Produce interior/casewrap files and order first six proofs.","Close payment, refund, deletion, photo-off, analytics, and legal-copy gaps while proofs run.","Complete final proofs, courier tests, live payment tests, and six dry runs before enabling the first paid block."]: num(x)

doc.add_heading("Sources",1)
for x in ["Google Cloud Service Terms — https://cloud.google.com/terms/service-terms","OpenAI Services Agreement — https://cdn.openai.com/osa/openai-services-agreement.pdf","MeitY DPDP commencement notification — https://www.meity.gov.in/static/uploads/2025/11/c56ceae6c383460ca69577428d36828b.pdf","CBIC GST rate schedule — https://cbic-gst.gov.in/hindi/gst-goods-services-rates.html","Consumer Protection E-Commerce Rules — https://consumeraffairs.nic.in/sites/default/files/E%20commerce%20rules_0.pdf","Razorpay checkout — https://razorpay.com/docs/developer-tools/integrations/standard-checkout/","Razorpay webhook validation — https://razorpay.com/docs/webhooks/validate-test/","CERT-In directions — https://www.cert-in.org.in/PDF/CERT-In_Directions_70B_28.04.2022.pdf","Baymard sales UX — https://baymard.com/blog/10-sales-ux-best-practices","Baymard checkout UX — https://baymard.com/blog/current-state-of-checkout-ux"]: bullet(x)

doc.save(OUT)
print(OUT)

